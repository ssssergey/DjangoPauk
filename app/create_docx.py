# -*- coding: utf-8 -*-
from django.http import HttpResponse
from docx import Document
from cStringIO import StringIO
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import os
document = Document()
def generate_docx(request):
    # document = Document()
    # u('тут добавляем в документ все что там нужно как в доке на python-docx')
    for i in range(3):
        local_docx(u'5 мая',u'Название',u'Текст')

    f = StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=test_result.docx'
    response['Content-Length'] = length
    return response


def local_docx(art_time, art_title, art_body):
    general_style = document.styles['Normal']
    font = general_style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    article_date = document.add_paragraph(art_time)
    article_date_format = article_date.paragraph_format
    article_date_format.first_line_indent = Cm(1.25)
    article_date_format.line_spacing = 1
    article_date_format.space_before = Pt(0)
    article_date_format.space_after = Pt(0)
    article_date_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    article_title = document.add_paragraph('')
    article_title.add_run(art_title).bold = True
    article_title_format = article_title.paragraph_format
    article_title_format.first_line_indent = Cm(1.25)
    article_title_format.line_spacing = 1
    article_title_format.space_before = Pt(0)
    article_title_format.space_after = Pt(0)
    article_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    article_body = document.add_paragraph(art_body)
    article_body_format = article_body.paragraph_format
    article_body_format.first_line_indent = Cm(1.25)
    article_body_format.line_spacing = 1
    article_body_format.space_before = Pt(0)
    article_body_format.space_after = Pt(0)
    article_body_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    document.add_paragraph('')

    # document.save(os.path.join(path, 'demo.docx'))


if __name__ == "__main__":
    path = "c:/Users/ASUS/Desktop/"
    document = Document()
    for i in range(3):
        local_docx(u'5 мая',u'Название',u'Текст')
    document.save(os.path.join(path, 'demo.docx'))