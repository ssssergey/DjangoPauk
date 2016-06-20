# -*- coding: utf-8 -*-
from django.shortcuts import render, render_to_response, get_object_or_404
from django.template import Context, RequestContext
from django.template.loader import get_template
from django.http import HttpResponse, Http404, HttpResponseRedirect
from django.contrib import auth, messages
from django.contrib.auth.models import User

from .models import News, Countries, UserCountry

from docx import Document
from cStringIO import StringIO
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from datetime import datetime, date, time, timedelta

# Create your views here.
def index(request):
    # User authentication
    uid = request.COOKIES.get('uid')
    if uid:
        try:
            user = User.objects.get(pk = int(uid))
        except User.DoesNotExist:
            user = create_new_user()
            uid = user.id
    else:
        user = create_new_user()
        uid = user.id
    user.backend='django.contrib.auth.backends.ModelBackend'
    auth.login(request, user)
    # countries = Countries.objects.order_by("name")
    response = render_to_response('index.html', {'user': user}, context_instance=RequestContext(request))
    response.set_cookie('uid', uid, max_age=60*60*24*365)
    # countries = Countries.objects.order_by("name")
    return response

# def list_by_country(request, slug):
#     # country = Countries.objects.get(slug=slug)
#     country = get_object_or_404(Countries, slug=slug)
#     news = News.objects.filter(country=country.name).order_by("-pub_time")[:100]
#     return render_to_response('app_list_by_country.html', locals())
#
# def details(request, news_id):
#     # news = News.objects.get(id=news_id)
#     news = get_object_or_404(News, id=news_id)
#     return render_to_response('app_details.html', locals())

def login(request):
    if request.method == 'POST':
        username = request.POST.get('username','')
        password = request.POST.get('password','')
        user = auth.authenticate(username=username, password=password)
        if user and user.is_active:
            auth.login(request, user)
            return HttpResponseRedirect('/')
        else:
            return HttpResponse(u'Неправильные данные')
    return render_to_response("login.html", context=RequestContext(request))


def create_new_user():
    '''
    Helper function
    :return: user object
    '''
    user = User.objects.create_user(username=u"Пользователь")
    user.save()
    user.username = u"Пользователь-{}".format(user.id)
    user.save()
    all_countries = Countries.objects.all()
    for country in all_countries:
        a = UserCountry(user=user, country=country, last_time=datetime.combine(date.today(), time()))
        a.save()
    return user

# def get_str_date_time(datetime_format):
#     month_names = ['января','февраля','марта','апреля','мая','июня','июля','августа','сентября','октября','ноября',
#                    'декабря']
#     month_name = month_names[datetime_format.month - 1]
#     date_final = '{} {} {} г.'.format(str(datetime_format.day).lstrip("0"),month_name,str(datetime_format.year))
#     time_final = datetime_format.strftime("%H.%M")
#     return date_final, time_final

def generate_docx(request):
    '''
    Helper function
    :param request:
    :return: docx document
    '''
    # selected_checkboxes = request.POST.getlist('ch_country')
    country = request.GET['name']
    slug = request.GET['slug']
    country_obj = Countries.objects.get(name=country)
    # slug = Countries.objects.get(name=country).slug
    UC_obj = UserCountry.objects.get(user=request.user, country=country_obj)
    last_article_time = UC_obj.last_time
    print last_article_time

    q = News.objects.all()
    q = q.filter(pub_time__gt = datetime.combine(date.today(), time()))
    q = q.filter(download_time__gt = last_article_time)
    q = q.filter(country=country).order_by('pub_time')

    if q.count() == 0:
        messages.info(request, u'Со времени последнего скачивания там НИЧЕГО НЕ СЛУЧИЛОСЬ. '
                               u'Попробуйте попозже.')
                      # .format(last_article_time.strftime('%H.%M')))
        return HttpResponseRedirect('/bd/select_country')

    try:
        UC_obj.last_time = q.order_by('-download_time')[0].download_time
        UC_obj.save()
    except IndexError:
        pass

    document = Document()
    general_style = document.styles['Normal']
    font = general_style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    sections = document.sections
    for section in sections:
        section.right_margin = Cm(2)

    content = document.add_paragraph(u'СОДЕРЖАНИЕ:')
    format = content.paragraph_format
    format.line_spacing = 1
    format.space_before = Pt(0)
    format.space_after = Pt(0)
    format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = content.add_run()
    run.add_break()

    for item in q:
        text = u' {} ("{}" <{}>)'.format(item.title, item.rss, item.pub_time.strftime("%H.%M"))
        content = document.add_paragraph(text, style='ListNumber')
        format = content.paragraph_format
        format.line_spacing = 1
        # format.space_before = Pt(4)
        # format.space_after = Pt(4)
        format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    document.add_paragraph('')

    for item in q:
        ttime = item.pub_time.strftime("%H.%M")
        print item.pub_time
        print item.download_time
        ddate = item.pub_time.strftime("%d.%m.%Y ") + u"г."
        rrss = item.rss
        ttitle = item.title
        bbody = item.body
        llink = item.link

        article_date = document.add_paragraph(u'({} ИА "{}")'.format(ttime, rrss))
        article_date_format = article_date.paragraph_format
        article_date_format.first_line_indent = Cm(1.25)
        article_date_format.line_spacing = 1
        article_date_format.space_before = Pt(0)
        article_date_format.space_after = Pt(0)
        article_date_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        article_title = document.add_paragraph('')
        article_title.add_run(ttitle).bold = True
        article_title_format = article_title.paragraph_format
        article_title_format.first_line_indent = Cm(1.25)
        article_title_format.line_spacing = 1
        article_title_format.space_before = Pt(0)
        article_title_format.space_after = Pt(0)
        article_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        for par in bbody.split('\n'):
            article_body = document.add_paragraph(par)
            article_body_format = article_body.paragraph_format
            article_body_format.first_line_indent = Cm(1.25)
            article_body_format.line_spacing = 1
            article_body_format.space_before = Pt(0)
            article_body_format.space_after = Pt(0)
            article_body_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # article_body = document.add_paragraph(llink)
        # article_body_format = article_body.paragraph_format
        # article_body_format.first_line_indent = Cm(1.25)
        # article_body_format.line_spacing = 1
        # article_body_format.space_before = Pt(0)
        # article_body_format.space_after = Pt(0)
        # article_body_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        article_body = document.add_paragraph(ddate)
        article_body_format = article_body.paragraph_format
        article_body_format.first_line_indent = Cm(1.25)
        article_body_format.line_spacing = 1
        article_body_format.space_before = Pt(0)
        article_body_format.space_after = Pt(0)
        article_body_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        document.add_paragraph('')

    f = StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=pauk_{:_<10}_{}.docx'.format(slug, datetime.now().strftime('%H.%M_%d.%m.%Y'))
    response['Content-Length'] = length

    return response