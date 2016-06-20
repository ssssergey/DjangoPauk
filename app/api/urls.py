# -*- coding: utf-8 -*-
from django.conf.urls import url
from django.http import HttpResponse, HttpResponseNotModified, StreamingHttpResponse
from wsgiref.util import FileWrapper
from django.contrib.auth.models import User
from rest_framework.serializers import ModelSerializer, HyperlinkedIdentityField
from rest_framework.response import Response
from rest_framework.generics import ListAPIView, RetrieveAPIView
from app.models import News, Countries, UserCountry

# Permissions
from rest_framework.permissions import AllowAny, IsAuthenticated, IsAdminUser, IsAuthenticatedOrReadOnly
# Filters
from rest_framework.filters import SearchFilter, OrderingFilter
# Pagination
from rest_framework.pagination import LimitOffsetPagination, PageNumberPagination

from rest_framework.decorators import api_view
from docx import Document
from cStringIO import StringIO
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime, date, time



class NewsLimitOffsetPagination(LimitOffsetPagination):
    max_limit = 500
    default_limit = 20

# Renderer
from rest_framework.renderers import JSONRenderer


class UTF8CharsetJSONRenderer(JSONRenderer):
    charset = 'utf-8'


# Serializers define the API representation.
class NewsListSerializer(ModelSerializer):
    url = HyperlinkedIdentityField(
        view_name='api:detail',
        lookup_field='id'
    )

    class Meta:
        model = News
        fields = ('id', 'url', 'rss', 'title', 'pub_time', 'download_time', 'country', 'link')


class NewsDetailSerializer(ModelSerializer):
    class Meta:
        model = News
        fields = ('id', 'rss', 'title', 'body', 'pub_time', 'download_time', 'country', 'link')

class GetCountriesSerializer(ModelSerializer):
    class Meta:
        model = Countries
        fields = ('slug', 'name')


# ViewSets define the view behavior.
class NewsListAPIView(ListAPIView):
    queryset = News.objects.all().order_by('-pub_time')
    serializer_class = NewsListSerializer
    filter_backends = [SearchFilter, OrderingFilter]
    search_fields = ['title', ]
    pagination_class = NewsLimitOffsetPagination
    # renderer_classes = [UTF8CharsetJSONRenderer]

class NewsCountryAPIView(ListAPIView):
    serializer_class = NewsListSerializer
    filter_backends = [SearchFilter, OrderingFilter]
    search_fields = ['title', ]
    pagination_class = NewsLimitOffsetPagination
    def get_queryset(self):
        country = self.kwargs['country']
        # result = News.objects.filter(country=country)
        result = Countries.objects.get(slug=country).news_set.all()
        return result


class NewsDetailAPIView(RetrieveAPIView):
    queryset = News.objects.all()
    serializer_class = NewsDetailSerializer
    # permission_classes = [IsAuthenticated]
    lookup_field = 'id'
    # renderer_classes = [UTF8CharsetJSONRenderer]

class GetCountriesAPIView(ListAPIView):
    queryset = Countries.objects.all().order_by("name")
    serializer_class = GetCountriesSerializer

@api_view()
def generate_docx(request):
    # selected_checkboxes = request.POST.getlist('ch_country')
    country = request.GET['name']
    slug = request.GET['slug']
    country_obj = Countries.objects.get(name=country)
    # slug = Countries.objects.get(name=country).slug
    UC_obj = UserCountry.objects.get(user=request.user, country=country_obj)
    last_article_time = UC_obj.last_time
    print last_article_time

    q = News.objects.all()
    q = q.filter(pub_time__gt = datetime.combine(date.today(), time()))
    q = q.filter(download_time__gt = last_article_time)
    q = q.filter(country=country).order_by('pub_time')

    if q.count() == 0:
        response = HttpResponseNotModified()
        response['Count'] = q.count()
        response['Content-Disposition'] = 'attachment; filename=empty'
        return response

    try:
        UC_obj.last_time = q.order_by('-download_time')[0].download_time
        UC_obj.save()
    except IndexError:
        pass

    document = Document()
    general_style = document.styles['Normal']
    font = general_style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    sections = document.sections
    for section in sections:
        section.right_margin = Cm(2)

    content = document.add_paragraph(u'СОДЕРЖАНИЕ:')
    format = content.paragraph_format
    format.line_spacing = 1
    format.space_before = Pt(0)
    format.space_after = Pt(0)
    format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = content.add_run()
    run.add_break()

    for item in q:
        text = u' {} ("{}" <{}>)'.format(item.title, item.rss, item.pub_time.strftime("%H.%M"))
        content = document.add_paragraph(text, style='ListNumber')
        format = content.paragraph_format
        format.line_spacing = 1
        # format.space_before = Pt(4)
        # format.space_after = Pt(4)
        format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    document.add_paragraph('')

    for item in q:
        ttime = item.pub_time.strftime("%H.%M")
        print item.pub_time
        print item.download_time
        ddate = item.pub_time.strftime("%d.%m.%Y ") + u"г."
        rrss = item.rss
        ttitle = item.title
        bbody = item.body
        llink = item.link

        article_date = document.add_paragraph(u'({} ИА "{}")'.format(ttime, rrss))
        article_date_format = article_date.paragraph_format
        article_date_format.first_line_indent = Cm(1.25)
        article_date_format.line_spacing = 1
        article_date_format.space_before = Pt(0)
        article_date_format.space_after = Pt(0)
        article_date_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        article_title = document.add_paragraph('')
        article_title.add_run(ttitle).bold = True
        article_title_format = article_title.paragraph_format
        article_title_format.first_line_indent = Cm(1.25)
        article_title_format.line_spacing = 1
        article_title_format.space_before = Pt(0)
        article_title_format.space_after = Pt(0)
        article_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        for par in bbody.split('\n'):
            article_body = document.add_paragraph(par)
            article_body_format = article_body.paragraph_format
            article_body_format.first_line_indent = Cm(1.25)
            article_body_format.line_spacing = 1
            article_body_format.space_before = Pt(0)
            article_body_format.space_after = Pt(0)
            article_body_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # article_body = document.add_paragraph(llink)
        # article_body_format = article_body.paragraph_format
        # article_body_format.first_line_indent = Cm(1.25)
        # article_body_format.line_spacing = 1
        # article_body_format.space_before = Pt(0)
        # article_body_format.space_after = Pt(0)
        # article_body_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        article_body = document.add_paragraph(ddate)
        article_body_format = article_body.paragraph_format
        article_body_format.first_line_indent = Cm(1.25)
        article_body_format.line_spacing = 1
        article_body_format.space_before = Pt(0)
        article_body_format.space_after = Pt(0)
        article_body_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        document.add_paragraph('')

    f = StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=pauk_{:_<10}_{}.docx'.format(slug, datetime.now().strftime('%H.%M_%d.%m.%Y'))
    response['Content-Length'] = length
    response['Count'] = q.count()
    return response

# Wire up our API using automatic URL routing.
# Additionally, we include login URLs for the browsable API.
urlpatterns = [
    url(r'(?P<id>[\d-]+)/$', NewsDetailAPIView.as_view(), name="detail"),
    url(r'countries/$', GetCountriesAPIView.as_view(), name="get_countries"),
    url(r'countries/(?P<country>.+)/$', NewsCountryAPIView.as_view(), name="country"),
    url(r'^$', NewsListAPIView.as_view(), name='list'),
    url(r'generate_docx/$', generate_docx, name='generate_docx')
]
