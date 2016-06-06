# -*- coding: utf-8 -*-
from django.shortcuts import render, render_to_response
from django.template import RequestContext
from django.http import HttpResponseRedirect, HttpResponse
from django.contrib.auth.models import User
from django.contrib import auth, messages
from app.models import Countries, UserCountry, News

from docx import Document
from cStringIO import StringIO
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from datetime import datetime, date, time, timedelta

# Create your views here.
def select_country(request):
    uid = request.COOKIES.get('uid')
    if uid:
        try:
            user = User.objects.get(pk = int(uid))
        except User.DoesNotExist:
            user = create_new_user()
            uid = user.id
    else:
        user = create_new_user()
        uid = user.id
    user.backend='django.contrib.auth.backends.ModelBackend'
    auth.login(request, user)
    countries = Countries.objects.order_by("name")
    response = render_to_response('bd_select_country.html', locals(), context_instance=RequestContext(request))
    response.set_cookie('uid', uid, max_age=60*60*24*365)
    return response



def create_new_user():
    user = User.objects.create_user(username=u"Пользователь")
    user.save()
    user.username = u"Пользователь-{}".format(user.id)
    user.save()
    all_countries = Countries.objects.all()
    for country in all_countries:
        a = UserCountry(user=user, country=country, last_time=datetime.combine(date.today(), time()))
        a.save()
    return user

# def get_str_date_time(datetime_format):
#     month_names = ['января','февраля','марта','апреля','мая','июня','июля','августа','сентября','октября','ноября',
#                    'декабря']
#     month_name = month_names[datetime_format.month - 1]
#     date_final = '{} {} {} г.'.format(str(datetime_format.day).lstrip("0"),month_name,str(datetime_format.year))
#     time_final = datetime_format.strftime("%H.%M")
#     return date_final, time_final

def generate_docx(request):
    # selected_checkboxes = request.POST.getlist('ch_country')
    country = request.GET['name']
    slug = request.GET['slug']
    country_obj = Countries.objects.get(name=country)
    # slug = Countries.objects.get(name=country).slug
    UC_obj = UserCountry.objects.get(user=request.user, country=country_obj)
    last_time_downloaded = UC_obj.last_time
    print last_time_downloaded

    q = News.objects.all()
    q = q.filter(pub_time__gt = datetime.combine(date.today(), time()))
    q = q.filter(download_time__gt = last_time_downloaded)
    q = q.filter(country=country).order_by('pub_time')

    UC_obj.last_time = datetime.now()
    UC_obj.save()

    if q.count() == 0:
        messages.info(request, u'ТАМ ПОКА НИЧЕГО НЕ СЛУЧИЛОСЬ. Попробуйте попозже.')
        return HttpResponseRedirect('/bd/select_country')

    document = Document()
    general_style = document.styles['Normal']
    font = general_style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    section = document.sections
    section.right_margin = Cm(1)

    content = document.add_paragraph(u'СОДЕРЖАНИЕ:')
    format = content.paragraph_format
    format.line_spacing = 1
    format.space_before = Pt(0)
    format.space_after = Pt(0)
    format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = content.add_run()
    run.add_break()

    for item in q:
        text = u'{} ("{}")'.format(item.title, item.rss)
        content = document.add_paragraph(text, style='ListNumber')
        format = content.paragraph_format
        format.line_spacing = 1
        format.space_before = Pt(0)
        format.space_after = Pt(0)
    document.add_paragraph('')

    for item in q:
        ttime = item.pub_time.strftime("%H.%M")
        print item.pub_time
        print item.download_time
        ddate = item.pub_time.strftime("%d.%m.%Y ") + u"г."
        rrss = item.rss
        ttitle = item.title
        bbody = item.body
        llink = item.link

        article_date = document.add_paragraph(u'({} ИА "{}")'.format(ttime, rrss))
        article_date_format = article_date.paragraph_format
        article_date_format.first_line_indent = Cm(1.25)
        article_date_format.line_spacing = 1
        article_date_format.space_before = Pt(0)
        article_date_format.space_after = Pt(0)
        article_date_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        article_title = document.add_paragraph('')
        article_title.add_run(ttitle).bold = True
        article_title_format = article_title.paragraph_format
        article_title_format.first_line_indent = Cm(1.25)
        article_title_format.line_spacing = 1
        article_title_format.space_before = Pt(0)
        article_title_format.space_after = Pt(0)
        article_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        for par in bbody.split('\n'):
            article_body = document.add_paragraph(par)
            article_body_format = article_body.paragraph_format
            article_body_format.first_line_indent = Cm(1.25)
            article_body_format.line_spacing = 1
            article_body_format.space_before = Pt(0)
            article_body_format.space_after = Pt(0)
            article_body_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # article_body = document.add_paragraph(llink)
        # article_body_format = article_body.paragraph_format
        # article_body_format.first_line_indent = Cm(1.25)
        # article_body_format.line_spacing = 1
        # article_body_format.space_before = Pt(0)
        # article_body_format.space_after = Pt(0)
        # article_body_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        article_body = document.add_paragraph(ddate)
        article_body_format = article_body.paragraph_format
        article_body_format.first_line_indent = Cm(1.25)
        article_body_format.line_spacing = 1
        article_body_format.space_before = Pt(0)
        article_body_format.space_after = Pt(0)
        article_body_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        document.add_paragraph('')

    f = StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=pauk_{:_<10}_{}.docx'.format(slug, datetime.now().strftime('%H.%M_%d.%m.%Y'))
    response['Content-Length'] = length

    return response