# -*- coding: utf-8 -*-
from app.models import News, Countries, UserCountry
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime, date, time
from django.contrib.auth.models import User
from django.contrib import auth
from django.shortcuts import render_to_response


def get_news_set(request, slug):
    '''
    Helper function. Returns queryset of fresh news and fixes datetime of last query on the country
    :param request:
    :param country:
    :return: q
    '''
    country_obj = Countries.objects.get(slug=slug)
    # slug = Countries.objects.get(name=country).slug
    UC_obj = UserCountry.objects.get(user=request.user, country=country_obj)
    last_article_time = UC_obj.last_time
    q = News.objects.all()
    q = q.filter(pub_time__gt=datetime.combine(date.today(), time()))
    q = q.filter(download_time__gt=last_article_time).order_by('pub_time')
    q = q.filter(country__slug=slug)
    return q, UC_obj


def create_docx(q):
    document = Document()
    general_style = document.styles['Normal']
    font = general_style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    sections = document.sections
    for section in sections:
        section.right_margin = Cm(2)

    content = document.add_paragraph(u'СОДЕРЖАНИЕ:')
    format = content.paragraph_format
    format.line_spacing = 1
    format.space_before = Pt(0)
    format.space_after = Pt(0)
    format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = content.add_run()
    run.add_break()

    for item in q:
        text = u' {} ("{}" <{}>)'.format(item.title, item.rss, item.pub_time.strftime("%H.%M"))
        content = document.add_paragraph(text, style='ListNumber')
        format = content.paragraph_format
        format.line_spacing = 1
        # format.space_before = Pt(4)
        # format.space_after = Pt(4)
        format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    document.add_paragraph('')

    for item in q:
        ttime = item.pub_time.strftime("%H.%M")
        print item.pub_time
        print item.download_time
        ddate = item.pub_time.strftime("%d.%m.%Y ") + u"г."
        rrss = item.rss
        ttitle = item.title
        bbody = item.body
        llink = item.link

        article_date = document.add_paragraph(u'({} ИА "{}")'.format(ttime, rrss))
        article_date_format = article_date.paragraph_format
        article_date_format.first_line_indent = Cm(1.25)
        article_date_format.line_spacing = 1
        article_date_format.space_before = Pt(0)
        article_date_format.space_after = Pt(0)
        article_date_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        article_title = document.add_paragraph('')
        article_title.add_run(ttitle).bold = True
        article_title_format = article_title.paragraph_format
        article_title_format.first_line_indent = Cm(1.25)
        article_title_format.line_spacing = 1
        article_title_format.space_before = Pt(0)
        article_title_format.space_after = Pt(0)
        article_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        for par in bbody.split('\n'):
            article_body = document.add_paragraph(par)
            article_body_format = article_body.paragraph_format
            article_body_format.first_line_indent = Cm(1.25)
            article_body_format.line_spacing = 1
            article_body_format.space_before = Pt(0)
            article_body_format.space_after = Pt(0)
            article_body_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # article_body = document.add_paragraph(llink)
        # article_body_format = article_body.paragraph_format
        # article_body_format.first_line_indent = Cm(1.25)
        # article_body_format.line_spacing = 1
        # article_body_format.space_before = Pt(0)
        # article_body_format.space_after = Pt(0)
        # article_body_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        article_body = document.add_paragraph(ddate)
        article_body_format = article_body.paragraph_format
        article_body_format.first_line_indent = Cm(1.25)
        article_body_format.line_spacing = 1
        article_body_format.space_before = Pt(0)
        article_body_format.space_after = Pt(0)
        article_body_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        document.add_paragraph('')
    return document


def create_new_user():
    '''
    Helper function
    :return: user object
    '''
    user = User.objects.create_user(username=u"Пользователь")
    user.save()
    user.username = u"Пользователь-{}".format(user.id)
    user.save()
    all_countries = Countries.objects.all()
    for country in all_countries:
        a = UserCountry(user=user, country=country, last_time=datetime.combine(date.today(), time()))
        a.save()
    return user


def auto_auth_new_user(request):
    # User authentication
    uid = request.COOKIES.get('uid')
    if uid:
        try:
            user = User.objects.get(pk=int(uid))
        except User.DoesNotExist:
            user = create_new_user()
            uid = user.id
    else:
        user = create_new_user()
        uid = user.id
    user.backend = 'django.contrib.auth.backends.ModelBackend'
    auth.login(request, user)
    return user, uid
